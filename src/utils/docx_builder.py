"""
DOCX Resume Builder

Takes a tailored resume dict (output from tailoring agent) and produces
a clean, ATS-friendly .docx file.

ATS rules:
- No tables, columns, text boxes, or graphics
- Standard section headers
- Single-column layout
- Consistent fonts (Calibri)
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

OUTPUT_DIR = Path(__file__).parent.parent.parent / "data" / "resumes"
OUTPUT_DIR.mkdir(exist_ok=True)

# Design constants
FONT_NAME = "Calibri"
COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)
COLOR_DARK_GRAY = RGBColor(0x40, 0x40, 0x40)
COLOR_MID_GRAY = RGBColor(0x60, 0x60, 0x60)

NAME_SIZE = Pt(20)
CONTACT_SIZE = Pt(10)
SECTION_HEADER_SIZE = Pt(11)
BODY_SIZE = Pt(10)
BULLET_SIZE = Pt(10)


def _set_font(run, size, bold=False, color=None):
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def _add_horizontal_rule(doc):
    """Add a thin horizontal line by setting a paragraph border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_section_header(doc, text):
    """Add a bold, spaced section header with underline rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    _set_font(run, SECTION_HEADER_SIZE, bold=True, color=COLOR_BLACK)

    # Bottom border on the paragraph acts as the section divider
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_contact_header(doc, contact: dict):
    """Add name and contact info at the top."""
    # Name
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_before = Pt(0)
    name_p.paragraph_format.space_after = Pt(2)
    run = name_p.add_run(contact.get("name", ""))
    _set_font(run, NAME_SIZE, bold=True, color=COLOR_BLACK)

    # Title (if present)
    title = contact.get("title", "")
    if title:
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_before = Pt(0)
        title_p.paragraph_format.space_after = Pt(3)
        run = title_p.add_run(title)
        _set_font(run, CONTACT_SIZE, bold=False, color=COLOR_DARK_GRAY)

    # Contact line: phone | email | linkedin | website | location
    parts = []
    if contact.get("phone"):
        parts.append(contact["phone"])
    if contact.get("email"):
        parts.append(contact["email"])
    if contact.get("linkedin"):
        parts.append(contact["linkedin"])
    if contact.get("website"):
        parts.append(contact["website"])
    if contact.get("location"):
        parts.append(contact["location"])

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_before = Pt(0)
    contact_p.paragraph_format.space_after = Pt(4)
    run = contact_p.add_run("  |  ".join(parts))
    _set_font(run, CONTACT_SIZE, color=COLOR_DARK_GRAY)


def _add_summary(doc, summary: str):
    """Add the professional summary section."""
    _add_section_header(doc, "Summary")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(summary.strip())
    _set_font(run, BODY_SIZE, color=COLOR_DARK_GRAY)


def _add_experience(doc, experience: list):
    """Add the experience section."""
    _add_section_header(doc, "Experience")

    for role in experience:
        # Role title + company on one line, dates right-aligned on same line
        role_p = doc.add_paragraph()
        role_p.paragraph_format.space_before = Pt(5)
        role_p.paragraph_format.space_after = Pt(1)

        title_run = role_p.add_run(role.get("title", ""))
        _set_font(title_run, BODY_SIZE, bold=True, color=COLOR_BLACK)

        company = role.get("company", "")
        if company:
            sep_run = role_p.add_run(f"  —  {company}")
            _set_font(sep_run, BODY_SIZE, bold=False, color=COLOR_DARK_GRAY)

        # Dates — right-aligned using a tab stop trick via a right-aligned run
        dates = role.get("dates", "")
        location = role.get("location", "")
        date_location = " | ".join(filter(None, [dates, location]))
        if date_location:
            # Use a tab to push dates to the right
            role_p.add_run("\t")
            dates_run = role_p.add_run(date_location)
            _set_font(dates_run, CONTACT_SIZE, color=COLOR_MID_GRAY)
            # Set tab stop at right margin
            pPr = role_p._p.get_or_add_pPr()
            tabs = OxmlElement("w:tabs")
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), "right")
            tab.set(qn("w:pos"), "9360")  # 6.5 inches in twips
            tabs.append(tab)
            pPr.append(tabs)

        # Bullets
        for bullet in role.get("bullets", []):
            bullet_p = doc.add_paragraph(style="List Bullet")
            bullet_p.paragraph_format.space_before = Pt(1)
            bullet_p.paragraph_format.space_after = Pt(1)
            bullet_p.paragraph_format.left_indent = Inches(0.2)
            run = bullet_p.add_run(bullet.strip())
            _set_font(run, BULLET_SIZE, color=COLOR_DARK_GRAY)


def _add_skills(doc, skills):
    """Add skills section as a comma-separated list."""
    _add_section_header(doc, "Skills")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)

    if isinstance(skills, list):
        skills_text = "  •  ".join(skills)
    else:
        skills_text = str(skills)

    run = p.add_run(skills_text)
    _set_font(run, BODY_SIZE, color=COLOR_DARK_GRAY)


def _set_margins(doc):
    """Set page margins to 0.7 inches all around for more content space."""
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)


def build_docx(tailored_resume: dict, output_filename: str) -> Path:
    """
    Build an ATS-friendly .docx from a tailored resume dict.

    Args:
        tailored_resume: The tailored_resume dict from the tailoring agent.
        output_filename: Filename (without extension) for the output file.

    Returns:
        Path to the saved .docx file.
    """
    doc = Document()
    _set_margins(doc)

    # Remove default paragraph spacing
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    # Clear any default paragraph the doc starts with
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    contact = tailored_resume.get("contact", {})
    summary = tailored_resume.get("summary", "")
    experience = tailored_resume.get("experience", [])
    skills = tailored_resume.get("skills", tailored_resume.get("skills_and_keywords", []))

    _add_contact_header(doc, contact)

    if summary:
        _add_summary(doc, summary)

    if experience:
        _add_experience(doc, experience)

    if skills:
        _add_skills(doc, skills)

    # Sanitize filename
    safe_name = re.sub(r"[^\w\-_]", "_", output_filename)
    output_path = OUTPUT_DIR / f"{safe_name}.docx"
    doc.save(output_path)

    return output_path
